import logging
import os
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter

logger = logging.getLogger(__name__)

# chunk_size=512, overlap~15% — benchmark Vecta 2026 đạt 69% accuracy tổng quát
_RECURSIVE_DEFAULTS = {"chunk_size": 512, "chunk_overlap": 75}


def chunk_semantic(docs: List[Document]) -> List[Document]:
    """
    Dùng cho tài liệu văn xuôi dài: báo cáo, hợp đồng, luận văn.
    Tách theo ranh giới ngữ nghĩa dựa trên embeddings.
    Lưu ý: tốn gấp đôi chi phí embedding so với recursive.
    """
    from langchain_experimental.text_splitter import SemanticChunker
    from langchain_openai import OpenAIEmbeddings

    splitter = SemanticChunker(
        OpenAIEmbeddings(model="text-embedding-3-small"),
        breakpoint_threshold_type="percentile",
        breakpoint_threshold_amount=95,
        # min_chunk_size tránh fragment quá ngắn làm giảm chất lượng generation
        min_chunk_size=150,
    )
    return splitter.split_documents(docs)


def chunk_recursive(
    docs: List[Document],
    chunk_size: int = _RECURSIVE_DEFAULTS["chunk_size"],
    chunk_overlap: int = _RECURSIVE_DEFAULTS["chunk_overlap"],
) -> List[Document]:
    """
    Dùng cho tài liệu tổng quát: sách, bài báo, tài liệu kỹ thuật, tiếng Việt.
    Chia theo ký tự phân cách tự nhiên (đoạn văn → câu → từ).
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    return splitter.split_documents(docs)


def chunk_fixed(
    docs: List[Document],
    chunk_size: int = 300,
    chunk_overlap: int = 50,
) -> List[Document]:
    """
    Dùng cho tài liệu có cấu trúc rõ ràng: FAQ, quy trình, bảng biểu.
    Chia theo kích thước cố định.
    """
    splitter = CharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separator="\n",
    )
    return splitter.split_documents(docs)


def _split_docx_by_heading(file_path: str, base_metadata: dict) -> List[Document]:
    """Đọc DOCX và tách theo Heading style, trả về List[Document] theo từng section."""
    from docx import Document as DocxDoc

    docx = DocxDoc(file_path)
    sections: List[Document] = []
    current_heading = "Mở đầu"
    current_lines: List[str] = []

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current_lines:
                sections.append(Document(
                    page_content="\n".join(current_lines),
                    metadata={**base_metadata, "section": current_heading},
                ))
            current_heading = text
            current_lines = []
        else:
            current_lines.append(text)

    if current_lines:
        sections.append(Document(
            page_content="\n".join(current_lines),
            metadata={**base_metadata, "section": current_heading},
        ))

    return [s for s in sections if s.page_content.strip()]


def chunk_document_based(docs: List[Document]) -> List[Document]:
    """
    Chia theo cấu trúc logic của tài liệu (Heading → section).
    Đạt ~87% accuracy với DOCX có heading rõ ràng (MDPI Bioengineering 2025).
    Tự động fallback sang recursive cho PDF hoặc DOCX không có heading.
    Yêu cầu: file gốc còn tồn tại tại doc.metadata['source'].
    """
    fallback_splitter = RecursiveCharacterTextSplitter(**_RECURSIVE_DEFAULTS)
    result: List[Document] = []

    for doc in docs:
        source = doc.metadata.get("source", "")
        file_type = doc.metadata.get("file_type", "")
        is_docx = file_type == "docx" or source.lower().endswith((".docx", ".doc"))

        if is_docx and source and os.path.exists(source):
            sections = _split_docx_by_heading(source, doc.metadata)
            if sections:
                # Các section quá dài vẫn cần chia thêm
                for section in sections:
                    if len(section.page_content) > _RECURSIVE_DEFAULTS["chunk_size"]:
                        result.extend(fallback_splitter.split_documents([section]))
                    else:
                        result.append(section)
                continue
        elif is_docx and source and not os.path.exists(source):
            logger.warning(
                "chunk_document_based: file không tồn tại tại '%s', fallback sang recursive. "
                "Gọi split_documents() TRƯỚC khi xóa temp file.",
                source,
            )

        # Fallback: PDF hoặc DOCX không có heading / không tìm thấy file gốc
        result.extend(fallback_splitter.split_documents([doc]))

    return result


STRATEGY_MAP = {
    "document": chunk_document_based,
    "recursive": chunk_recursive,
    "semantic": chunk_semantic,
    "fixed": chunk_fixed,
}

STRATEGY_LABELS = {
    "document": "Theo cấu trúc (Document) — DOCX có heading, tài liệu pháp lý",
    "recursive": "Đệ quy (Recursive) — tài liệu tổng quát, tiếng Việt",
    "semantic": "Ngữ nghĩa (Semantic) — báo cáo dài, hợp đồng",
    "fixed": "Cố định (Fixed) — FAQ, quy trình, bảng biểu",
}


def split_documents(docs: List[Document], strategy: str = "recursive") -> List[Document]:
    if strategy not in STRATEGY_MAP:
        raise ValueError(f"Chiến lược không hợp lệ: '{strategy}'. Chọn một trong: {list(STRATEGY_MAP.keys())}")
    return STRATEGY_MAP[strategy](docs)
